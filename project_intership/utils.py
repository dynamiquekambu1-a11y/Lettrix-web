import os
from datetime import datetime
from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.lib.utils import ImageReader


# ===============================
#        PATHS
# ===============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

EXPORT_WORD_DIR = os.path.join(BASE_DIR, "export", "word")
EXPORT_PDF_DIR = os.path.join(BASE_DIR, "export", "pdf")

os.makedirs(EXPORT_WORD_DIR, exist_ok=True)
os.makedirs(EXPORT_PDF_DIR, exist_ok=True)


# ===============================
#        UTIL
# ===============================
def make_filename(prefix, ext):
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}_{ts}.{ext}"


# ===============================
#        WORD
# ===============================
def export_to_word(text, meta):
    path = os.path.join(
        EXPORT_WORD_DIR, make_filename("internship", "docx")
    )

    doc = Document()
    doc.add_heading("INTERNSHIP APPLICATION LETTER", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(path)
    return path


# ===============================
#        PDF
# ===============================
def export_to_pdf(text, meta):
    path = os.path.join(
        EXPORT_PDF_DIR, make_filename("internship", "pdf")
    )

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=7.5 * cm,
        bottomMargin=3 * cm,
    )

    # -------- Styles --------
    body_style = ParagraphStyle(
        "body", fontSize=11, leading=15, alignment=TA_JUSTIFY
    )

    title_style = ParagraphStyle(
        "title",
        fontSize=16,
        alignment=TA_CENTER,
        textColor=HexColor("#1f3c88"),
        fontName="Helvetica-Bold",
        spaceAfter=14,
        spaceBefore=0,


    )

    signature_style = ParagraphStyle(
        "signature",
        fontSize=11,
        alignment=TA_LEFT,
    )

    footer_style = ParagraphStyle(
        "footer",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=HexColor("#777777"),
    )

    story = []

    # ---- LETTER TITLE (PLACED CORRECTLY) ----
    story.append(Spacer(1, 5))
    story.append(
        Paragraph("INTERNSHIP APPLICATION LETTER", title_style)
    )

    # ---- LETTER BODY ----
    for p in text.split("\n\n"):
        if p.strip():
            story.append(
                Paragraph(p.replace("\n", "<br/>"), body_style)
            )
            story.append(Spacer(1, 10))

    # ---- SIGNATURE ----
    story.append(Spacer(1, 30))
    story.append(
        Paragraph(
            "<b>Signature :</b><br/><br/>"
            "..............................................................",
            signature_style,
        )
    )

    # ---- FOOTER ----
    story.append(Spacer(1, 40))
    story.append(
        Paragraph(
            "This letter was generated professionally using LETTRIX – WEB.",
            footer_style,
        )
    )

    # ===============================
    #        HEADER + WATERMARK
    # ===============================
    def draw_header(canvas, doc):
        canvas.saveState()
        w, h = A4
        top_y = h - 2.5 * cm

        # ---- WATERMARK ----
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 60)
        canvas.setFillColor(HexColor("#E6E6E6"))
        canvas.translate(w / 2, h / 2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, "LETTRIX – WEB")
        canvas.restoreState()

        # ---- LOGO ----
        logo_path = meta.get("company_logo_path")
        if logo_path and os.path.isfile(logo_path):
            try:
                logo = ImageReader(logo_path)
                iw, ih = logo.getSize()
                lw = 3 * cm
                lh = lw * ih / iw
                canvas.drawImage(
                    logo,
                    2 * cm,
                    top_y - lh + 2 * cm,  # ⬆️ logo légèrement plus haut
                    lw,
                    lh,
                    mask="auto",
                )

            except Exception:
                pass

        # ---- COMPANY NAME (OPTIONAL) ----
        if meta.get("company_name"):
            canvas.setFont("Helvetica-Bold", 15)
            canvas.setFillColor(HexColor("#1f3c88"))
            canvas.drawCentredString(
                w / 2, top_y, meta.get("company_name")
            )

        # ---- ADDRESS ----
        if meta.get("company_address"):
            canvas.setFont("Helvetica", 10)
            canvas.setFillColor(HexColor("#444444"))
            canvas.drawCentredString(
                w / 2, top_y - 18, meta.get("company_address")
            )

        # ---- BLUE LINE ----
        canvas.setStrokeColor(HexColor("#1f3c88"))
        canvas.setLineWidth(1.5)
        canvas.line(
            2 * cm, top_y - 32, w - 2 * cm, top_y - 32
        )

        # ---- RECIPIENT INFO ----
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(HexColor("#F4B400"))
        RIGHT_y = top_y - 55

        if meta.get("email"):
            canvas.drawString(
                2 * cm, RIGHT_y, meta.get("email")
            )
            RIGHT_y -= 12

        if meta.get("phone"):
            canvas.drawString(
                2 * cm, RIGHT_y, meta.get("phone")
            )

        canvas.restoreState()

    doc.build(
        story,
        onFirstPage=draw_header,
        onLaterPages=draw_header,
    )

    return path


# ===============================
#        VALIDATION
# ===============================
def validate_required_fields(data, required):
    return [k for k in required if not data.get(k)]

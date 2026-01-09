import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, Color
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER
from reportlab.lib.units import mm

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ==============================
#        PATHS
# ==============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, ".."))

EXPORT_PDF = os.path.join(PROJECT_ROOT, "export", "pdf")
EXPORT_WORD = os.path.join(PROJECT_ROOT, "export", "word")

os.makedirs(EXPORT_PDF, exist_ok=True)
os.makedirs(EXPORT_WORD, exist_ok=True)


# ==============================
#        COLORS
# ==============================
BLUE = HexColor("#0d6efd")
YELLOW = HexColor("#FFD700")
BLACK = HexColor("#000000")
WATERMARK = Color(0.6, 0.6, 0.6, alpha=0.15)


# ==============================
#        WATERMARK
# ==============================
def draw_watermark(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 60)
    canvas.setFillColor(WATERMARK)
    canvas.translate(300, 400)
    canvas.rotate(45)
    canvas.drawCentredString(0, 0, "RESIGNATION LETTER")
    canvas.restoreState()


# ==============================
#        PDF EXPORT
# ==============================
def export_to_pdf(text, metadata=None):
    metadata = metadata or {}
    path = os.path.join(EXPORT_PDF, "resignation_letter.pdf")

    if os.path.exists(path):
        os.remove(path)

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=25 * mm,
        leftMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    elements = []

    # -------- METADATA --------
    company = metadata.get("company_name", "")
    email = metadata.get("company_email", "")
    phone = metadata.get("company_phone", "")

    # -------- LOGO (UPLOAD USER) --------
    logo = ""
    logo_url = metadata.get("company_logo", "").strip()

    if logo_url:
        # convert URL -> disk path
        logo_url = logo_url.replace("/projects/", "").lstrip("/")
        logo_path = os.path.join(PROJECT_ROOT, logo_url)

        if os.path.isfile(logo_path):
            logo = Image(logo_path, 40, 40)

    # -------- HEADER --------
    header = Table(
        [
            [
                Paragraph(
                    f"<b><font color='#0d6efd'>{company}</font></b>",
                    ParagraphStyle("h", fontSize=14)
                ),
                logo,
            ],
            [
                Paragraph(
                    f"<font color='black'>{email} | {phone}</font>",
                    ParagraphStyle("e", fontSize=9)
                ),
                "",
            ],
        ],
        colWidths=[120 * mm, 40 * mm],
    )

    header.setStyle(TableStyle([
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))

    elements.append(header)

    # -------- BLUE LINE --------
    elements.append(
        Table([[""]], colWidths=[170 * mm],
              style=[("LINEBELOW", (0, 0), (-1, -1), 2, BLUE)])
    )
    elements.append(Spacer(1, 15))

    # -------- TITLE --------
    elements.append(Paragraph(
        "<b><font color='#0d6efd'>RESIGNATION LETTER</font></b>",
        ParagraphStyle("title", fontSize=16, alignment=TA_CENTER)
    ))
    elements.append(Spacer(1, 20))

    # -------- BODY --------
    body = ParagraphStyle("body", fontSize=11, leading=15)
    for line in text.split("\n"):
        if line.strip():
            elements.append(Paragraph(line, body))
            elements.append(Spacer(1, 6))

    # -------- SIGNATURE --------
    elements.append(Spacer(1, 40))
    elements.append(Paragraph(
        "<font color='#FFD700'><b>Signature : _________________________</b></font>",
        ParagraphStyle("sig", alignment=TA_RIGHT)
    ))

    doc.build(elements, onFirstPage=draw_watermark, onLaterPages=draw_watermark)
    return path


# ==============================
#        WORD EXPORT
# ==============================
def export_to_word(text, metadata=None):
    metadata = metadata or {}
    path = os.path.join(EXPORT_WORD, "resignation_letter.docx")

    if os.path.exists(path):
        os.remove(path)

    doc = Document()

    # -------- HEADER --------
    p = doc.add_paragraph(metadata.get("company_name", ""))
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True

    p = doc.add_paragraph(
        f"{metadata.get('company_email','')} | {metadata.get('company_phone','')}"
    )
    p.runs[0].font.size = Pt(9)

    # -------- TITLE --------
    p = doc.add_paragraph("\nRESIGNATION LETTER")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True

    # -------- BODY --------
    for line in text.split("\n"):
        doc.add_paragraph(line)

    # -------- SIGNATURE --------
    p = doc.add_paragraph("\nSignature : _________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.bold = True

    doc.save(path)
    return path

# outil_web.py
import os
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.units import mm


# ==============================
#  Helper : extraire nom de famille
# ==============================
def extract_surname(fullname: str) -> str:
    if not fullname:
        return ""
    fullname = fullname.strip()
    fullname = fullname.splitlines()[0].strip()
    parts = fullname.split()
    return parts[-1] if parts else fullname


# ==============================
#       WORD GENERATOR
# ==============================
def generer_word_temp(lettre: dict) -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    path = tmp.name
    tmp.close()

    doc = Document()

    # ----------------------- Expéditeur -----------------------
    exp_nom = lettre.get("exp_nom") or lettre.get("nom") or ""
    exp_adresse = lettre.get("exp_adresse", "")
    exp_cp = lettre.get("exp_cp", "")
    exp_ville = lettre.get("exp_ville", "")
    exp_email = lettre.get("exp_email", "")
    exp_tel = lettre.get("exp_tel", "")

    p = doc.add_paragraph()
    p.add_run(f"{exp_nom}\n{exp_adresse}\n{exp_cp} {exp_ville}\n{exp_email} | {exp_tel}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(14)

    # ----------------------- Destinataire -----------------------
    destinataire = lettre.get("destinataire", "")
    nom_entreprise = lettre.get("nom_entreprise", "")
    dest_adresse = lettre.get("dest_adresse", "")
    dest_cp = lettre.get("dest_cp", "")
    dest_ville = lettre.get("dest_ville", "")

    p = doc.add_paragraph()
    p.add_run(f"{destinataire}\n{nom_entreprise}\n{dest_adresse}\n{dest_cp} {dest_ville}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.paragraph_format.space_after = Pt(14)

    # ----------------------- Objet -----------------------
    objet = lettre.get("objet", "").strip()
    if objet:
        p = doc.add_paragraph()
        r = p.add_run(f"Objet : {objet}")
        r.bold = True
        r.underline = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(18)

    # ----------------------- Civilité -----------------------
    surname = extract_surname(destinataire)
    civ = f"Monsieur / Madame {surname}," if surname else "Monsieur / Madame,"

    p = doc.add_paragraph()
    p.add_run(civ)
    p.paragraph_format.space_after = Pt(14)

    # ----------------------- Alignements -----------------------
    align_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        None: WD_PARAGRAPH_ALIGNMENT.LEFT,
        "": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }

    intro_align = lettre.get("intro_align", "left")
    corps_align = lettre.get("corps_align", "left")
    conc_align = lettre.get("conc_align", "left")

    # ----------------------- Texte (intro / corps / conc) -----------------------
    intro = lettre.get("introduction", "").strip()
    corps = lettre.get("corps", "").strip()
    conc = lettre.get("conclusion", "").strip()

    if intro:
        p = doc.add_paragraph(intro)
        p.alignment = align_map[intro_align]
        p.paragraph_format.space_after = Pt(14)

    if corps:
        p = doc.add_paragraph(corps)
        p.alignment = align_map[corps_align]
        p.paragraph_format.space_after = Pt(14)

    if conc:
        p = doc.add_paragraph(conc)
        p.alignment = align_map[conc_align]
        p.paragraph_format.space_after = Pt(18)

    # ----------------------- Date -----------------------
    p = doc.add_paragraph(f"Fait à {exp_ville}, le {datetime.now().strftime('%d/%m/%Y')}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.paragraph_format.space_after = Pt(100)

    # ----------------------- Signature -----------------------
    signature = lettre.get("signature", "").strip()
    if signature:
        p = doc.add_paragraph(signature)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_after = Pt(900)

    doc.save(path)
    return path


# ==============================
#          PDF GENERATOR
# ==============================
def generer_pdf_temp(lettre: dict) -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    path = tmp.name
    tmp.close()

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm
    )

    elements = []

    # Styles
    style = ParagraphStyle(name="Normal", fontName="Times-Roman", fontSize=11, leading=16)
    style_right = ParagraphStyle(name="Right", parent=style, alignment=TA_RIGHT)
    style_center = ParagraphStyle(name="Center", parent=style, alignment=TA_CENTER)

    # ----------------------- Expéditeur -----------------------
    exp_nom = lettre.get("exp_nom") or lettre.get("nom") or ""
    exp_adresse = lettre.get("exp_adresse", "")
    exp_cp = lettre.get("exp_cp", "")
    exp_ville = lettre.get("exp_ville", "")
    exp_email = lettre.get("exp_email", "")
    exp_tel = lettre.get("exp_tel", "")

    elements.append(Paragraph(f"{exp_nom}<br/>{exp_adresse}<br/>{exp_cp} {exp_ville}<br/>{exp_email} | {exp_tel}", style))
    elements.append(Spacer(1, 12))

    # ----------------------- Destinataire -----------------------
    destinataire = lettre.get("destinataire", "")
    nom_entreprise = lettre.get("nom_entreprise", "")
    dest_adresse = lettre.get("dest_adresse", "")
    dest_cp = lettre.get("dest_cp", "")
    dest_ville = lettre.get("dest_ville", "")

    elements.append(Paragraph(f"{destinataire}<br/>{nom_entreprise}<br/>{dest_adresse}<br/>{dest_cp} {dest_ville}", style_right))
    elements.append(Spacer(1, 12))

    # ----------------------- Objet -----------------------
    objet = lettre.get("objet", "").strip()
    if objet:
        elements.append(Paragraph(f"<b><u>Objet : {objet}</u></b>", style_center))
        elements.append(Spacer(1, 12))

    # ----------------------- Civilité -----------------------
    surname = extract_surname(destinataire)
    civ = f"Monsieur / Madame {surname}," if surname else "Monsieur / Madame,"
    elements.append(Paragraph(civ, style))
    elements.append(Spacer(1, 12))

    # ----------------------- Alignements -----------------------
    align_map = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT}

    intro_align = lettre.get("intro_align", "left")
    corps_align = lettre.get("corps_align", "left")
    conc_align = lettre.get("conc_align", "left")

    intro = lettre.get("introduction", "").strip()
    corps = lettre.get("corps", "").strip()
    conc = lettre.get("conclusion", "").strip()

    if intro:
        st = ParagraphStyle(name="Intro", parent=style, alignment=align_map[intro_align])
        elements.append(Paragraph(intro, st))
        elements.append(Spacer(1, 12))

    if corps:
        st = ParagraphStyle(name="Corps", parent=style, alignment=align_map[corps_align])
        elements.append(Paragraph(corps, st))
        elements.append(Spacer(1, 12))

    if conc:
        st = ParagraphStyle(name="Conc", parent=style, alignment=align_map[conc_align])
        elements.append(Paragraph(conc, st))
        elements.append(Spacer(1, 12))

    # ----------------------- Date -----------------------
    date_style = ParagraphStyle(name="Date", parent=style, alignment=TA_RIGHT)
    elements.append(Spacer(1, 80))
    elements.append(Paragraph(f"Fait à {exp_ville}, le {datetime.now().strftime('%d/%m/%Y')}", date_style))
    elements.append(Spacer(1, 20))

    # ----------------------- Signature -----------------------
    signature = lettre.get("signature", "").strip()
    if signature:
        elements.append(Paragraph(signature, style))

    doc.build(elements)
    return path

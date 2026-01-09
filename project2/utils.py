import os, datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.pagesizes import A4

import os, datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.pagesizes import A4
import os, datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.pagesizes import A4
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

EXPORT_WORD_DIR = os.path.join(BASE_DIR, "export", "word")
EXPORT_PDF_DIR = os.path.join(BASE_DIR, "export", "pdf")

os.makedirs(EXPORT_WORD_DIR, exist_ok=True)
os.makedirs(EXPORT_PDF_DIR, exist_ok=True)

def make_filename(prefix="document", ext="pdf"):
    ts = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}_{ts}.{ext}"

def export_to_word_leave(text, metadata=None):
    doc = Document()
    doc.add_heading("LEAVE REQUEST LETTER", level=1)
    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)
    filename = make_filename("leave_request","docx")
    path = os.path.join(EXPORT_WORD_DIR, filename)
    doc.save(path)
    return path


def export_leave_pdf(text, metadata=None):
    metadata = metadata or {}

    filename = f"leave_request_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    path = os.path.join(EXPORT_PDF_DIR, filename)

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    story = []

    # ========= STYLES (NOMS UNIQUES) =========
    company_header_style = ParagraphStyle(
        "CompanyHeaderStyle",
        fontSize=15,
        leading=18,
        alignment=TA_CENTER,
        textColor=HexColor("#f4b400"),  # JAUNE OK
        spaceBefore=0,
        spaceAfter=4
    )

    company_address_style = ParagraphStyle(
        "CompanyAddressStyle",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=HexColor("#444444"),
        spaceAfter=8
    )

    body_style = ParagraphStyle(
        "BodyStyle",
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=10
    )

    footer_style = ParagraphStyle(
        "FooterStyle",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=HexColor("#888888"),
        spaceBefore=12
    )

    # ========= LOGO =========
    logo_web_path = metadata.get("company_logo", "").strip()

    if logo_web_path:
        # Transformer le chemin web en chemin disque réel
        if logo_web_path.startswith("/"):
            logo_disk_path = os.path.join(BASE_DIR, logo_web_path.lstrip("/"))
        else:
            logo_disk_path = os.path.join(BASE_DIR, logo_web_path)

        logo_disk_path = os.path.abspath(logo_disk_path)

        if os.path.exists(logo_disk_path):
            logo = Image(logo_disk_path, width=3 * cm, height=3 * cm)
            logo.hAlign = "RIGHT"
            story.append(logo)
            story.append(Spacer(1, 12))
        else:
            print("LOGO NOT FOUND:", logo_disk_path)  # debug

    # ========= HEADER VISUEL =========
    company_name = metadata.get("company_name", "").strip()
    company_address = metadata.get("company_address", "").strip()

    if company_name:
        story.append(Paragraph(company_name, company_header_style))

    if company_address:
        story.append(Paragraph(company_address, company_address_style))

    story.append(
        HRFlowable(
            width="100%",
            thickness=1.5,
            color=HexColor("#f4b400")  # JAUNE GARANTI
        )
    )

    story.append(Spacer(1, 16))

    # ========= CONTENU =========
    for p in text.split("\n\n"):
        p = p.strip()
        if p:
            story.append(Paragraph(p.replace("\n", "<br/>"), body_style))

    # ========= FOOTER =========
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc")))
    story.append(Paragraph(
        "Generated with LETTRIX – WEB",
        footer_style
    ))

    # ========= WATERMARK =========
    def watermark(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 55)
        canvas.setFillColor(Color(0.6, 0.6, 0.6, alpha=0.12))
        canvas.translate(A4[0]/2, A4[1]/2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, "LETTRIX – WEB")
        canvas.restoreState()

    doc.build(story, onFirstPage=watermark, onLaterPages=watermark)
    return path


def validate_required_fields(data: dict, required=None):
    required = required or []
    missing = []
    for k in required:
        if not data.get(k) or not str(data.get(k)).strip():
            missing.append(k)
    return missing

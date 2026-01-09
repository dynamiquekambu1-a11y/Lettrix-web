import os
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image,
    HRFlowable, Table, TableStyle
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.units import cm
from reportlab.lib.colors import black, HexColor, Color
from reportlab.lib.utils import ImageReader


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
EXPORT_WORD_DIR = os.path.join(BASE_DIR, "export", "word")
EXPORT_PDF_DIR = os.path.join(BASE_DIR, "export", "pdf")

os.makedirs(EXPORT_WORD_DIR, exist_ok=True)
os.makedirs(EXPORT_PDF_DIR, exist_ok=True)


def make_filename(prefix, ext):
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}_{ts}.{ext}"


# ---------------- WORD ----------------
def export_to_word(text, metadata=None):
    metadata = metadata or {}
    applicant = metadata.get("applicant_name", "applicant").replace(" ", "_")
    path = os.path.join(EXPORT_WORD_DIR, f"{applicant}_Job_Application.docx")

    doc = Document()
    doc.add_heading("JOB APPLICATION LETTER", level=1)

    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.save(path)
    return path

def _watermark(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 48)
    canvas.setFillColor(Color(0.7, 0.7, 0.7, alpha=0.12))
    canvas.translate(A4[0] / 2, A4[1] / 2)
    canvas.rotate(45)
    canvas.drawCentredString(0, 0, "LETTRIX")
    canvas.restoreState()


def export_to_pdf(text, metadata=None):
    metadata = metadata or {}

    logo_path = metadata.get("company_logo_path")
    company_name = metadata.get("company_name", "")
    company_address = metadata.get("company_address", "")
    company_email = metadata.get("company_email", "")
    company_phone = metadata.get("company_phone", "")
    applicant_name = metadata.get("applicant_name", "")

    filename = make_filename(applicant_name.replace(" ", "_") or "application", "pdf")
    path = os.path.join(EXPORT_PDF_DIR, filename)

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    story = []

    # -------- LOGO --------
    if logo_path and os.path.isfile(logo_path):
        img = ImageReader(logo_path)
        iw, ih = img.getSize()
        w = 3 * cm
        h = w * ih / iw
        story.append(Image(logo_path, width=w, height=h))
        story.append(Spacer(1, 12))

    # -------- STYLES --------
    title = ParagraphStyle(
        "title",
        fontSize=16,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        textColor=HexColor("#003366")
    )

    meta = ParagraphStyle(
        "meta",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=HexColor("#666666")
    )

    section_title = ParagraphStyle(
        "section_title",
        fontSize=13,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        textColor=HexColor("#003366")
    )

    body = ParagraphStyle(
        "body",
        fontSize=11,
        leading=15,
        alignment=TA_JUSTIFY
    )

    footer = ParagraphStyle(
        "footer",
        fontSize=8,
        alignment=TA_CENTER,
        textColor=HexColor("#888888")
    )

    sign_style = ParagraphStyle(
        "sign",
        fontSize=10,
        alignment=TA_RIGHT
    )

    # -------- HEADER --------
    story.append(Paragraph(company_name, title))
    story.append(Spacer(1, 8))

    story.append(Paragraph(
        "<br/>".join(filter(None, [
            company_address,
            f"Email: {company_email}" if company_email else "",
            f"Phone: {company_phone}" if company_phone else ""
        ])),
        meta
    ))

    # -------- BLUE LINE --------
    story.append(Spacer(1, 18))  # ✅ PLUS D’ESPACE
    story.append(HRFlowable(
        width="100%",
        thickness=1.5,
        color=HexColor("#003366")
    ))
    story.append(Spacer(1, 20))  # ✅ TEXTE DESCEND BIEN

    # -------- JOB APPLICATION TITLE --------
    story.append(Paragraph("JOB APPLICATION", section_title))
    story.append(Spacer(1, 18))

    # -------- LETTER BODY --------
    for p in text.split("\n\n"):
        story.append(Paragraph(p.replace("\n", "<br/>"), body))
        story.append(Spacer(1, 10))

    # -------- SIGNATURE --------
    story.append(Spacer(1, 35))
    story.append(Paragraph("Signature: .....................................", sign_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(applicant_name, ParagraphStyle(
        "sign_name",
        fontSize=10,
        alignment=TA_RIGHT,
        fontName="Helvetica-Bold"
    )))

    # -------- FOOTER TEXT --------
    story.append(Spacer(1, 40))
    story.append(Paragraph(
        "This letter was generated automatically by Lettrix-WE",
        footer
    ))

    doc.build(
        story,
        onFirstPage=_watermark,
        onLaterPages=_watermark
    )

    return path


def validate_required_fields(data, required):
    return [k for k in required if not data.get(k)]

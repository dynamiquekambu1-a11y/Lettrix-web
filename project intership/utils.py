from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.utils import ImageReader
import os
from datetime import datetime

EXPORT_WORD_DIR = "export/word"
EXPORT_PDF_DIR = "export/pdf"

def make_filename(prefix, ext):
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}_{ts}.{ext}"

# ------------------ WORD ------------------
def export_to_word(text, meta):
    os.makedirs(EXPORT_WORD_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading("INTERNSHIP APPLICATION LETTER", level=1)

    # Ajouter chaque ligne
    for line in text.split("\n"):
        doc.add_paragraph(line)

    path = os.path.join(EXPORT_WORD_DIR, make_filename("internship", "docx"))
    doc.save(path)
    return path

# ------------------ PDF ------------------
def export_to_pdf(text, meta):
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    os.makedirs(os.path.join(BASE_DIR, EXPORT_PDF_DIR), exist_ok=True)
    path = os.path.join(BASE_DIR, EXPORT_PDF_DIR, make_filename("internship", "pdf"))

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=4.8*cm,
        bottomMargin=3*cm
    )

    story = []

    # ---------- STYLES ----------
    title_style = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=14, textColor=HexColor("#1f3c88"), alignment=TA_CENTER, spaceAfter=14)
    normal_style = ParagraphStyle("normal", fontName="Helvetica", fontSize=11, leading=15, alignment=TA_JUSTIFY)
    meta_style = ParagraphStyle("meta", fontName="Helvetica", fontSize=10, textColor=HexColor("#555555"))
    footer_style = ParagraphStyle("footer", fontName="Helvetica", fontSize=8, textColor=HexColor("#777777"), alignment=TA_RIGHT)

    # ---------- CONTENU ----------
    story.append(Paragraph("INTERNSHIP APPLICATION LETTER", title_style))
    story.append(Paragraph(f"<b>Date:</b> {meta.get('date','')}", meta_style))
    story.append(Spacer(1, 16))

    for p in text.split("\n\n"):
        story.append(Paragraph(p.replace("\n", "<br/>"), normal_style))
        story.append(Spacer(1, 12))

    story.append(Spacer(1, 30))
    story.append(Paragraph("Signature: ____________________", meta_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph(meta.get("applicant_name", ""), footer_style))
    story.append(Spacer(1, 25))
    story.append(Paragraph("Generated professionally using LETTRIX – WEB.", footer_style))

    # ---------- HEADER ----------
    def draw_header(canvas, doc):
        canvas.saveState()
        page_width, page_height = A4

        # ----- Logo dynamique -----
        logo_path = meta.get("company_logo", "static/logo.png").replace("\\", "/")
        if not os.path.isabs(logo_path):
            logo_path = os.path.join(BASE_DIR, logo_path)
        logo_path = os.path.abspath(logo_path)

        print("LOGO PATH =", logo_path)
        print("LOGO EXISTS =", os.path.exists(logo_path))

        if os.path.exists(logo_path):
            try:
                logo_img = ImageReader(logo_path)
                iw, ih = logo_img.getSize()
                logo_width = 3 * cm
                logo_height = logo_width * (ih / iw)
                canvas.drawImage(
                    logo_img,
                    (page_width - logo_width) / 2,
                    page_height - 3.5 * cm,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=True,
                    mask="auto"
                )
            except Exception as e:
                print("Erreur logo:", e)

        # ----- Company name -----
        company_name = meta.get("company_name", "")
        if company_name:
            canvas.setFont("Helvetica-Bold", 12)
            canvas.setFillColor(HexColor("#1f3c88"))
            canvas.drawCentredString(page_width / 2, page_height - 4.1*cm, company_name)

        # ----- Ligne header -----
        canvas.setStrokeColor(HexColor("#1f3c88"))
        canvas.setLineWidth(0.6)
        canvas.line(2*cm, page_height - 4.3*cm, page_width - 2*cm, page_height - 4.3*cm)

        # ----- Watermark -----
        canvas.setFont("Helvetica-Bold", 60)
        canvas.setFillColor(Color(0.7, 0.7, 0.7, alpha=0.08))
        canvas.saveState()
        canvas.translate(page_width / 2, page_height / 2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, "LETTRIX")
        canvas.restoreState()

        canvas.restoreState()

    # ---------- FOOTER ----------
    def draw_footer(canvas, doc):
        canvas.saveState()
        page_width, _ = A4
        canvas.setFont("Helvetica-Bold", 10)
        canvas.setFillColor(HexColor("#1f3c88"))
        canvas.drawCentredString(page_width / 2, 1.8*cm, meta.get("company_name", ""))
        canvas.restoreState()

    doc.build(
        story,
        onFirstPage=lambda c,d: (draw_header(c,d), draw_footer(c,d)),
        onLaterPages=lambda c,d: (draw_header(c,d), draw_footer(c,d))
    )

    return path

# ------------------ VALIDATION ------------------
def validate_required_fields(data, required):
    return [k for k in required if not data.get(k)]

import os
import datetime

from docx import Document
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image,
    HRFlowable, Table, TableStyle
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.units import cm
from reportlab.lib.colors import black, HexColor, Color


# ==============================
#        PATHS GLOBAUX
# ==============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, ".."))

EXPORT_WORD_DIR = os.path.join(PROJECT_ROOT, "export", "word")
EXPORT_PDF_DIR = os.path.join(PROJECT_ROOT, "export", "pdf")

os.makedirs(EXPORT_WORD_DIR, exist_ok=True)
os.makedirs(EXPORT_PDF_DIR, exist_ok=True)


# ==============================
#        UTILITAIRES
# ==============================
def make_filename(prefix="certificate", ext="pdf"):
    ts = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}_{ts}.{ext}"


# ==============================
#        EXPORT WORD
# ==============================
def export_to_word(text, metadata=None):
    doc = Document()
    doc.add_heading("WORK CERTIFICATE", level=1)

    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)

    filename = make_filename(prefix="certificate", ext="docx")
    path = os.path.join(EXPORT_WORD_DIR, filename)
    doc.save(path)
    return path


# ==============================
#        EXPORT PDF
# ==============================
def export_to_pdf(text, metadata=None):
    metadata = metadata or {}

    # --------- METADATA ----------
    company_name = metadata.get("company_name", "")
    company_address = metadata.get("company_address", "")
    company_phone = metadata.get("company_phone", "")
    company_email = metadata.get("company_email", "")
    employee_name = metadata.get("employee_name", "")
    position = metadata.get("position", "")
    department = metadata.get("department", "")
    start_date = metadata.get("start_date", "")
    end_date = metadata.get("end_date", "")
    contract_type = metadata.get("contract_type", "")
    signer_name = metadata.get("signer_name", "")
    signer_role = metadata.get("signer_role", "")
    signature_place = metadata.get("signature_place", "")
    signature_date = metadata.get("signature_date", "")

    # 🔴 LOGO = CHEMIN FICHIER RÉEL (UPLOAD UTILISATEUR)
    logo_path = metadata.get("company_logo", "").strip()
    if logo_path and not os.path.isabs(logo_path):
        logo_path = os.path.normpath(os.path.join(PROJECT_ROOT, logo_path))

    # --------- PDF FILE ----------
    filename = make_filename(prefix="certificate", ext="pdf")
    path = os.path.join(EXPORT_PDF_DIR, filename)

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )

    story = []

    # --------- STYLES ----------
    title_style = ParagraphStyle(
        "Title",
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        textColor=HexColor("#0033cc"),
    )

    subtitle_style = ParagraphStyle(
        "Subtitle",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=HexColor("#555555"),
    )

    normal = ParagraphStyle(
        "Normal",
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
    )

    right_style = ParagraphStyle(
        "Right",
        fontSize=10,
        alignment=TA_RIGHT,
        fontName="Helvetica-Bold",
    )

    small_center = ParagraphStyle(
        "SmallCenter",
        fontSize=8,
        alignment=TA_CENTER,
        textColor=HexColor("#666666"),
    )

    # ==============================
    #        HEADER (LOGO UPLOAD)
    # ==============================
    if logo_path and os.path.isfile(logo_path):
        try:
            img = Image(logo_path)
            img.drawHeight = 3 * cm
            img.drawWidth = 3 * cm
            img.hAlign = "RIGHT"
            story.append(img)
            story.append(Spacer(1, 6))
        except Exception as e:
            print("LOGO ERROR:", e)

    if company_name:
        story.append(Paragraph(company_name, title_style))

    header_lines = []
    if company_address:
        header_lines.append(company_address)

    contacts = []
    if company_phone:
        contacts.append(f"Phone: {company_phone}")
    if company_email:
        contacts.append(f"Email: {company_email}")

    if contacts:
        header_lines.append(" | ".join(contacts))

    if header_lines:
        story.append(Paragraph("<br/>".join(header_lines), subtitle_style))

    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=black))
    story.append(Spacer(1, 12))

    # ==============================
    #        TITLE
    # ==============================
    story.append(Paragraph("WORK CERTIFICATE", title_style))
    story.append(Spacer(1, 12))

    # ==============================
    #        EMPLOYEE TABLE
    # ==============================
    rows = []
    if employee_name:
        rows.append([
            Paragraph('<font color="#f1c40f">Employee:</font>', normal),
            employee_name
        ])

    if position:
        rows.append([Paragraph('<font color="#f1c40f">Position:</font>', normal), position])
    if department:
        rows.append([Paragraph('<font color="#f1c40f">Department:</font>', normal), department])
    if contract_type:
        rows.append([Paragraph('<font color="#f1c40f">Contract type:</font>', normal), contract_type])
    if start_date:
        rows.append([Paragraph('<font color="#f1c40f">Start date:</font>', normal), start_date])
    if end_date:
        rows.append([Paragraph('<font color="#f1c40f">End date:</font>', normal), end_date])

    if rows:
        table = Table(rows, colWidths=[90, 350])
        table.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
        story.append(Spacer(1, 12))

    # ==============================
    #        TEXT
    # ==============================
    story.append(Paragraph("To whom it may concern,", normal))
    story.append(Spacer(1, 6))

    for paragraph in text.split("\n\n"):
        if paragraph.strip():
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), normal))
            story.append(Spacer(1, 6))

    story.append(Spacer(1, 18))

    # ==============================
    #        SIGNATURE
    # ==============================
    if signer_name:
        story.append(Paragraph(signer_name, right_style))
    if signer_role:
        story.append(Paragraph(signer_role, right_style))
    if signature_place or signature_date:
        story.append(
            Paragraph(f"{signature_place} {signature_date}".strip(), right_style)
        )

    story.append(Spacer(1, 20))
    story.append(
        Paragraph('<font color="#f1c40f">Signature: ____________________</font>', normal)
    )

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "This certificate is issued upon the request of the concerned employee.",
        small_center
    ))

    # ==============================
    #        WATERMARK
    # ==============================
    def add_watermark(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 70)
        canvas.setFillColor(Color(0.6, 0.6, 0.6, alpha=0.12))
        canvas.translate(300, 400)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, "LETTRIX - WEB")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_watermark, onLaterPages=add_watermark)

    return path


# ==============================
#        VALIDATION
# ==============================
def validate_required_fields(data: dict, required=None):
    required = required or []
    missing = []
    for k in required:
        if not data.get(k) or not str(data.get(k)).strip():
            missing.append(k)
    return missing
